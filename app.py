import streamlit as st
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import io
import re
from copy import deepcopy
from docx.oxml.ns import qn

# --- 1. LOGIKA DETEKSI OTOMATIS (HEURISTIC) ---
def auto_detect_sections(ms_doc):
    """
    Mendeteksi bagian naskah secara otomatis dengan fitur:
    1. Stop-logic pada Afiliasi agar tidak bocor ke Pendahuluan.
    2. Konversi tanda koma (,) ke titik koma (;) pada Kata Kunci.
    3. Pemeliharaan label formal untuk Email dan Abstrak.
    """
    sections = {
        "Judul": "", "Author": "", "Afiliasi": "", "Email": "",
        "Email Korespondensi": "", 
        "Abstrak": "", "Kata Kunci": "", "Abstract (EN)": "", "Keywords (EN)": ""
    }
    
    # Ambil semua paragraf yang tidak kosong
    paragraphs = [p.text.strip() for p in ms_doc.paragraphs if p.text.strip()]
    if not paragraphs:
        return sections

    # 1. Judul & Author (Posisi baris 1 & 2)
    sections["Judul"] = paragraphs[0]
    if len(paragraphs) > 1:
        sections["Author"] = paragraphs[1]

    import re
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    
    afiliasi_list = []
    found_abstract_start = False # Flag pengunci afiliasi

    # Scan mulai dari baris ke-3
    for i in range(2, len(paragraphs)):
        text = paragraphs[i]
        text_upper = text.upper()
        
        # --- A. DETEKSI MARKER ABSTRAK (Marker Berhenti untuk Afiliasi) ---
        if text_upper.startswith("ABSTRAK") or text_upper.startswith("ABSTRACT"):
            found_abstract_start = True 
            
            # Logika Abstrak Indonesia
            if text_upper.startswith("ABSTRAK") and not sections["Abstrak"]:
                content_buffer = []
                # Ambil teks setelah kata 'Abstrak' di baris yang sama
                first_line = text[7:].strip(" :-").strip()
                if first_line: content_buffer.append(first_line)
                
                # Ambil baris-baris berikutnya sampai ketemu 'Kata Kunci'
                for j in range(i + 1, len(paragraphs)):
                    if "KATA KUNCI" in paragraphs[j].upper(): break
                    content_buffer.append(paragraphs[j])
                
                # Gabungkan dengan label standar (akan di-bold di build_auto_docx)
                sections["Abstrak"] = f"Abstrak{' '.join(content_buffer)}"
            
            # Logika Abstract English
            elif text_upper.startswith("ABSTRACT") and not sections["Abstract (EN)"]:
                content_buffer = []
                first_line = text[8:].strip(" :-").strip()
                if first_line: content_buffer.append(first_line)
                
                for j in range(i + 1, len(paragraphs)):
                    if "KEYWORDS" in paragraphs[j].upper() or "PENDAHULUAN" in paragraphs[j].upper(): break
                    content_buffer.append(paragraphs[j])
                
                sections["Abstract (EN)"] = f"Abstract{' '.join(content_buffer)}"

        # --- B. LOGIKA KATA KUNCI & KEYWORDS (Standardisasi Titik Koma) ---
        elif "KATA KUNCI" in text_upper:
            val = text.split(":", 1)[-1].strip() if ":" in text else text.replace("KATA KUNCI", "").strip()
            # Bersihkan tanda baca dan ganti koma ke titik koma
            val = val.replace(",", ";")
            keywords_cleaned = "; ".join([k.strip() for k in val.split(";") if k.strip()])
            sections["Kata Kunci"] = f"Kata Kunci: {keywords_cleaned}"

        elif "KEYWORDS" in text_upper:
            val = text.split(":", 1)[-1].strip() if ":" in text else text.replace("KEYWORDS", "").strip()
            val = val.replace(",", ";")
            keywords_cleaned = "; ".join([k.strip() for k in val.split(";") if k.strip()])
            sections["Keywords (EN)"] = f"Keywords: {keywords_cleaned}"

        # --- C. LOGIKA EMAIL ---
        elif "@" in text:
            emails = re.findall(email_pattern, text)
            if emails:
                if any(k in text_upper for k in ["CORRESPONDING", "KORESPONDENSI", "*"]):
                    sections["Email Korespondensi"] = f"Email Penulis Korespondensi: {emails[0]}"
                elif not sections["Email"]:
                    sections["Email"] = f"Email: {', '.join(emails)}"
        
        # --- D. LOGIKA PENGUMPULAN AFILIASI ---
        else:
            # Ambil teks hanya jika gerbang Abstrak belum terbuka
            if not found_abstract_start and len(text) > 3:
                afiliasi_list.append(text)

    # Gabungkan semua baris afiliasi dengan baris baru (\n)
    sections["Afiliasi"] = "\n".join(afiliasi_list)
    
    return sections

# --- 2. FUNGSI CLONING (DIPERBAIKI) ---
def move_body_elements(source_doc, target_doc):
    import io
    import re
    from copy import deepcopy
    from docx.table import Table
    
    start_found = False
    source_body = source_doc.element.body
    target_body = target_doc.element.body

    # --- DEFINISI STYLE ---
    STYLE_SUBJUDUL_UTAMA = "Subjudul_Jurnal"  # Bab 1, 2, 4
    STYLE_HASIL_PEMBAHASAN = "Sub_Judul"      # Bab 3
    STYLE_SUBSUB = "Subsubjudul_Jurnal"       # 2.1, 3.1, dst
    STYLE_ISI = "Isi_Jurnal"                  # Paragraf isi & isi tabel

    # --- DAFTAR MARKER ---
    METODOLOGI_MARKERS = ["2. METODOLOGI PENELITIAN", "2. RESEARCH METHODOLOGY", "2. METODE PENELITIAN", "2. RESEARCH METHOD"]
    HASIL_MARKERS = ["3. HASIL DAN PEMBAHASAN", "3. RESULTS AND DISCUSSION", "3. HASIL PENELITIAN DAN PEMBAHASAN", "3. HASIL"]
    KESIMPULAN_MARKERS = ["4. KESIMPULAN", "4. CONCLUSION", "4. PENUTUP", "KESIMPULAN DAN SARAN"]
    
    subsub_pattern = r'^\d+\.\d+'

    # 1. Koleksi Binary Gambar
    image_blobs = {}
    for rel in source_doc.part.rels.values():
        if "image" in rel.target_ref:
            image_blobs[rel.rId] = rel.target_part.blob

    # 2. Iterasi Elemen
    for element in source_body.iterchildren():
        text = ""
        is_paragraph = element.tag.endswith('p')
        
        if is_paragraph:
            text_parts = [node.text for node in element.iter() if node.tag.endswith('t') and node.text is not None]
            text = "".join(text_parts).strip()
            text_upper = text.upper()
            
            has_image = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or \
                        element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            
            if not text and not has_image:
                continue 

        # Trigger Mulai dari Bab 1
        if not start_found and is_paragraph and ("1. PENDAHULUAN" in text_upper or "1. INTRODUCTION" in text_upper):
            start_found = True

        if start_found:
            # JALUR A: GAMBAR & CAPTION
            if is_paragraph and has_image:
                new_p = target_doc.add_paragraph()
                for run_element in element.iterchildren():
                    if run_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or \
                       run_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
                        for node in run_element.iter():
                            rid = node.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            if rid in image_blobs:
                                image_stream = io.BytesIO(image_blobs[rid])
                                new_p.add_run().add_picture(image_stream)
                    t_tags = run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    for t in t_tags:
                        if t.text:
                            new_p.add_run(t.text)

                apply_style_to_element(new_p._element, STYLE_ISI)
                target_body.append(new_p._element)

            # JALUR B: TABEL
            elif element.tag.endswith('tbl'):
                new_element = deepcopy(element)
                new_table = Table(new_element, target_doc)
                try:
                    new_table.style = 'Tabel_Jurnal'
                except:
                    new_table.style = 'Table Grid'
                new_table.autofit = True
                for row in new_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            apply_style_to_element(paragraph._element, STYLE_ISI)
                target_body.append(new_element)

            # JALUR C: TEKS BIASA (Termasuk Isi Kesimpulan)
            else:
                new_element = deepcopy(element)
                if is_paragraph:
                    # Cek Judul Bab 3
                    if any(marker in text_upper for marker in HASIL_MARKERS):
                        apply_style_to_element(new_element, STYLE_HASIL_PEMBAHASAN)
                    
                    # Cek Judul Bab 1, 2, atau 4
                    elif any(marker in text_upper for marker in (METODOLOGI_MARKERS + KESIMPULAN_MARKERS)) or \
                         "1. PENDAHULUAN" in text_upper or "1. INTRODUCTION" in text_upper:
                        apply_style_to_element(new_element, STYLE_SUBJUDUL_UTAMA)
                    
                    # Cek Sub-bab (3.1, 4.1, dsb)
                    elif re.match(subsub_pattern, text) and len(text) < 150:
                        apply_style_to_element(new_element, STYLE_SUBSUB)
                    
                    # SEMUA ISI (Pendahuluan, Metodologi, Pembahasan, KESIMPULAN)
                    else:
                        apply_style_to_element(new_element, STYLE_ISI)
                
                target_body.append(new_element)

def apply_style_to_element(element, style_name):
    """
    Suntik style dan bersihkan format 'sampah' (highlight, shading, color)
    agar benar-benar mengikuti template.
    """
    # 1. Terapkan Style Paragraf
    pPr = element.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = pPr.makeelement(qn('w:pStyle'))
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), style_name)

    # 2. Bersihkan Shading/Highlight di level Paragraf
    shd_p = pPr.find(qn('w:shd'))
    if shd_p is not None:
        pPr.remove(shd_p)

    # 3. Iterasi ke setiap Run (teks) untuk membersihkan highlight & shading
    for run in element.findall(qn('w:r')):
        rPr = run.get_or_add_rPr()
        
        # Hapus Highlight (warna stabilo)
        highlight = rPr.find(qn('w:highlight'))
        if highlight is not None:
            rPr.remove(highlight)
            
        # Hapus Shading (warna latar belakang/bayangan teks)
        shading = rPr.find(qn('w:shd'))
        if shading is not None:
            rPr.remove(shading)

        # Hapus Warna Teks (agar kembali otomatis mengikuti style template)
        color = rPr.find(qn('w:color'))
        if color is not None:
            rPr.remove(color)

def move_body_elements(source_doc, target_doc):
    import io
    import re
    from copy import deepcopy
    from docx.table import Table
    
    # Status awal: pencarian dimulai dari Pendahuluan
    start_found = False
    source_body = source_doc.element.body
    target_body = target_doc.element.body

    # --- KONFIGURASI NAMA STYLE (SESUAIKAN DENGAN TEMPLATE ANDA) ---
    STYLE_SUBJUDUL_UTAMA = "Subjudul_Jurnal"  # Untuk Bab 1, 2, 4
    STYLE_HASIL_PEMBAHASAN = "Sub_Judul"      # Khusus Bab 3
    STYLE_SUBSUB = "Subsubjudul_Jurnal"       # Untuk 2.1, 3.1, dst
    STYLE_ISI = "Isi_Jurnal"                  # Paragraf teks & isi tabel

    # --- DAFTAR MARKER DETEKSI BAB ---
    METODOLOGI_MARKERS = ["2. METODOLOGI PENELITIAN", "2. RESEARCH METHODOLOGY", "2. METODE PENELITIAN", "2. RESEARCH METHOD"]
    HASIL_MARKERS = ["3. HASIL DAN PEMBAHASAN", "3. RESULTS AND DISCUSSION", "3. HASIL PENELITIAN DAN PEMBAHASAN", "3. HASIL"]
    KESIMPULAN_MARKERS = ["4. KESIMPULAN", "4. CONCLUSION", "4. PENUTUP", "KESIMPULAN DAN SARAN"]
    
    subsub_pattern = r'^\d+\.\d+' # Pola untuk angka seperti 2.1 atau 3.2.1

    # --- 1. PRA-PEMROSESAN: KOLEKSI BINARY GAMBAR ---
    # Mengambil semua gambar dari naskah asli agar bisa disisipkan ulang secara aman
    image_blobs = {}
    for rel in source_doc.part.rels.values():
        if "image" in rel.target_ref:
            image_blobs[rel.rId] = rel.target_part.blob

    # --- 2. ITERASI SETIAP ELEMEN DI DALAM NASKAH ASLI ---
    for element in source_body.iterchildren():
        text = ""
        is_paragraph = element.tag.endswith('p')
        
        if is_paragraph:
            # Ambil semua teks di dalam paragraf tanpa merusak format internal
            text_parts = [node.text for node in element.iter() if node.tag.endswith('t') and node.text is not None]
            text = "".join(text_parts).strip()
            text_upper = text.upper()
            
            # Deteksi apakah paragraf ini mengandung objek gambar
            has_image = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or \
                        element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
            
            # Abaikan paragraf kosong yang tidak memiliki gambar
            if not text and not has_image:
                continue 

        # TRIGGER MULAI: Hanya proses elemen SETELAH menemukan Pendahuluan
        if not start_found and is_paragraph and ("1. PENDAHULUAN" in text_upper or "1. INTRODUCTION" in text_upper):
            start_found = True

        if start_found:
            # --- JALUR A: PARAGRAF DENGAN GAMBAR (RE-BUILD METHOD) ---
            if is_paragraph and has_image:
                new_p = target_doc.add_paragraph()
                
                # Iterasi anak elemen (Run) untuk memastikan urutan Gambar -> Teks Caption tetap benar
                for run_element in element.iterchildren():
                    # Jika elemen adalah Drawing (Gambar)
                    if run_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or \
                       run_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
                        
                        for node in run_element.iter():
                            rid = node.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            if rid in image_blobs:
                                image_stream = io.BytesIO(image_blobs[rid])
                                new_p.add_run().add_picture(image_stream)
                    
                    # Jika elemen adalah Text (Caption/Keterangan)
                    t_tags = run_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    for t in t_tags:
                        if t.text:
                            new_p.add_run(t.text)

                apply_style_to_element(new_p._element, STYLE_ISI)
                target_body.append(new_p._element)

            # --- JALUR B: TABEL ---
            elif element.tag.endswith('tbl'):
                new_element = deepcopy(element)
                new_table = Table(new_element, target_doc)
                
                # Terapkan Style Tabel dari Template (Pastikan nama style sesuai di Word)
                try:
                    new_table.style = 'Tabel_Jurnal'
                except:
                    new_table.style = 'Table Grid'
                
                new_table.autofit = True
                
                # Format font di dalam setiap sel agar konsisten dengan Isi Jurnal
                for row in new_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            apply_style_to_element(paragraph._element, STYLE_ISI)
                
                target_body.append(new_element)

            # --- JALUR C: TEKS BIASA (DEEPCOPY METHOD) ---
            else:
                new_element = deepcopy(element)
                if is_paragraph:
                    # 1. Cek jika Judul Bab 3 (Hasil & Pembahasan)
                    if any(marker in text_upper for marker in HASIL_MARKERS):
                        apply_style_to_element(new_element, STYLE_HASIL_PEMBAHASAN)
                    
                    # 2. Cek jika Judul Bab 2 atau Bab 4 (Metodologi / Kesimpulan)
                    elif any(marker in text_upper for marker in (METODOLOGI_MARKERS + KESIMPULAN_MARKERS)):
                        apply_style_to_element(new_element, STYLE_SUBJUDUL_UTAMA)
                    
                    # 3. Cek jika Sub-bab (Contoh: 3.1 Analisis Data)
                    elif re.match(subsub_pattern, text) and len(text) < 150:
                        apply_style_to_element(new_element, STYLE_SUBSUB)
                    
                    # 4. Paragraf isi standar
                    else:
                        apply_style_to_element(new_element, STYLE_ISI)
                
                target_body.append(new_element)
            
# --- 3. FUNGSI BUILDER ---
def build_auto_docx(template_file, manuscript_file, data_map):
    new_doc = Document(template_file)
    ms_doc = Document(manuscript_file)
    
    # 1. Bersihkan isi template asli
    for p in new_doc.paragraphs: p._element.getparent().remove(p._element)
    for t in new_doc.tables: t._element.getparent().remove(t._element)

    # 2. Definisi Pemetaan Style Otomatis sesuai standar Anda
    style_mapping = {
        "Judul": "Judul_Jurnal",
        "Author": "Author_Jurnal",
        "Afiliasi": "Afiliasi_Jurnal",
        "Email": "Afiliasi_Jurnal",
        "Email Korespondensi": "Afiliasi_Jurnal",
        "Abstrak": "Abstrak_Jurnal",
        "Kata Kunci": "Abstrak_Jurnal",
        "Abstract (EN)": "Abstrak_Jurnal",
        "Keywords (EN)": "Abstrak_Jurnal"
    }

    # Urutan output di dokumen
    order = ["Judul", "Author", "Afiliasi", "Email", "Email Korespondensi", 
             "Abstrak", "Kata Kunci", "Abstract (EN)", "Keywords (EN)"]

    # Label yang harus di-BOLD (hanya kata kuncinya saja)
    special_labels = {
        "Abstrak": "Abstrak",
        "Abstract (EN)": "Abstract",
        "Kata Kunci": "Kata Kunci:",
        "Keywords (EN)": "Keywords:"
    }

    for cat in order:
        content = data_map.get(cat, "").strip()
        if content:
            # Ambil nama style dari mapping, jika tidak ada gunakan 'Normal'
            target_style_name = style_mapping.get(cat, "Normal")
            p = new_doc.add_paragraph()
            
            try: 
                p.style = new_doc.styles[target_style_name]
            except: 
                # Jika style tidak ditemukan di template, fallback ke Normal
                p.style = new_doc.styles["Normal"]
            
            # Logika Penebalan Parsial
            if cat in special_labels:
                label_to_bold = special_labels[cat]
                if content.startswith(label_to_bold):
                    run_label = p.add_run(label_to_bold)
                    run_label.bold = True
                    
                    remaining_text = content[len(label_to_bold):]
                    run_content = p.add_run(remaining_text)
                    run_content.bold = None 
                else:
                    run = p.add_run(content)
                    run.bold = None
            else:
                run = p.add_run(content)
                run.bold = None

            # Reset Font agar mengikuti settingan Style di Template (Font Name, Size, Color, etc.)
            for run in p.runs:
                run.font.name = None
                run.font.size = None

    # 3. Masukkan Isi Utama (Pendahuluan dst.)
    move_body_elements(ms_doc, new_doc)

    buffer = io.BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. UI STREAMLIT ---
st.set_page_config(page_title="Auto Journal Formatter", layout="wide")
st.title("🚀 Fully Automated Journal Formatter")
st.markdown("Sistem mendeteksi bagian naskah dan menerapkan **Style Jurnal** secara otomatis.")

if 'detected_data' not in st.session_state:
    st.session_state.detected_data = None

# Mapping Style Otomatis (Sama dengan yang di Builder)
AUTO_STYLE_MAP = {
    "Judul": "Judul_Jurnal",
    "Author": "Author_Jurnal",
    "Afiliasi": "Afiliasi_Jurnal",
    "Email": "Afiliasi_Jurnal",
    "Email Korespondensi": "Afiliasi_Jurnal",
    "Abstrak": "Abstrak_Jurnal",
    "Kata Kunci": "Abstrak_Jurnal",
    "Abstract (EN)": "Abstrak_Jurnal",
    "Keywords (EN)": "Abstrak_Jurnal"
}

u1, u2 = st.columns(2)
with u1:
    tpl_file = st.file_uploader("📂 1. Upload Template Jurnal", type="docx")
with u2:
    ms_file = st.file_uploader("📝 2. Upload Naskah Mentah", type="docx")

if tpl_file and ms_file:
    if st.button("🔍 Deteksi Bagian Otomatis", use_container_width=True):
        doc_ms = Document(ms_file)
        st.session_state.detected_data = auto_detect_sections(doc_ms)

if st.session_state.detected_data:
    st.write("---")
    st.subheader("⚙️ Verifikasi Hasil Deteksi")
    st.info("Style telah diterapkan otomatis berdasarkan standar: Judul_Jurnal, Author_Jurnal, Afiliasi_Jurnal, dan Abstrak_Jurnal.")
    
    col_grid = st.columns(2)
    
    # Loop untuk menampilkan hasil deteksi saja (tanpa selectbox style)
    for i, (cat, val) in enumerate(st.session_state.detected_data.items()):
        with col_grid[i % 2]:
            # Update data jika pengguna melakukan pengeditan manual di text_area
            st.session_state.detected_data[cat] = st.text_area(
                f"Bagian: {cat} (Style: {AUTO_STYLE_MAP.get(cat)})", 
                val, 
                height=150, 
                key=f"in_{cat}"
            )

    if st.button("📥 Generate & Download Naskah", use_container_width=True):
        # Kita panggil builder dengan data terbaru
        # Builder sekarang hanya butuh tpl, ms, dan data_map (karena style sudah hardcoded di dalam builder)
        final_out = build_auto_docx(tpl_file, ms_file, st.session_state.detected_data)
        st.download_button(
            label="✅ Klik untuk Unduh Hasil Formating",
            data=final_out,
            file_name="Formatted_Journal_Final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )